"""DOCX/XLSX/PDF + safe file ops with backups."""
from __future__ import annotations
import hashlib
import shutil
import time
from pathlib import Path

from ..core.config import BACKUPS
from ..core.logging import get
from ..core.safety import gate, Action, Tier

log = get("omni.files")


def sha256(p: Path) -> str:
    h = hashlib.sha256()
    with open(p, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


async def backup(path: str | Path) -> Path:
    src = Path(path)
    if not src.exists():
        raise FileNotFoundError(src)
    BACKUPS.mkdir(parents=True, exist_ok=True)
    dst = BACKUPS / f"{src.stem}_{int(time.time())}{src.suffix}"
    shutil.copy2(src, dst)
    log.info("backup %s -> %s", src, dst)
    return dst


async def read_text(path: str | Path) -> str:
    p = Path(path)
    if not gate.folder_allowed(p.parent):
        await gate.gate(Action("file.read", Tier.ACT, {"path": str(p)}))
    suf = p.suffix.lower()
    if suf == ".docx":
        return _docx_text(p)
    if suf == ".pdf":
        return _pdf_text(p)
    if suf in {".xlsx", ".xlsm"}:
        return _xlsx_text(p)
    return p.read_text(encoding="utf-8", errors="ignore")


def _docx_text(p: Path) -> str:
    from docx import Document
    doc = Document(str(p))
    return "\n".join(par.text for par in doc.paragraphs)


def _pdf_text(p: Path) -> str:
    from pypdf import PdfReader
    r = PdfReader(str(p))
    return "\n".join((page.extract_text() or "") for page in r.pages)


def _xlsx_text(p: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(p), data_only=True, read_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            out.append("\t".join("" if v is None else str(v) for v in row))
    return "\n".join(out)


# DOCX edit
async def docx_replace(src: Path, dst: Path, replacements: dict[str, str]) -> Path:
    """Copy src -> dst, then run case-sensitive paragraph-level replacements."""
    await gate.gate(Action(
        "file.write", Tier.ACT,
        {"src": str(src), "dst": str(dst), "n": len(replacements)},
        suggested="Apply tailored resume edits to job-specific copy",
    ))
    await backup(src)
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)
    from docx import Document
    doc = Document(str(dst))
    for par in doc.paragraphs:
        for old, new in replacements.items():
            if old in par.text:
                for run in par.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
    doc.save(str(dst))
    return dst


# tracker (XLSX)
TRACKER_HEADERS = [
    "date", "company", "role", "location", "url", "source",
    "score", "verdict", "cover_letter_path", "resume_path",
    "status", "notes", "questions_asked", "answers_learned",
]


async def tracker_append(path: str | Path, row: dict) -> None:
    p = Path(path)
    await gate.gate(Action(
        "file.write", Tier.ACT, {"path": str(p)},
        suggested="Append row to job tracker",
    ))
    if p.exists():
        await backup(p)
    from openpyxl import Workbook, load_workbook
    if p.exists():
        wb = load_workbook(str(p))
        ws = wb.active
    else:
        p.parent.mkdir(parents=True, exist_ok=True)
        wb = Workbook()
        ws = wb.active
        ws.append(TRACKER_HEADERS)
    ws.append([row.get(h, "") for h in TRACKER_HEADERS])
    wb.save(str(p))
